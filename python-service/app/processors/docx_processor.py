from typing import Dict, List, Any, Optional
import docx
from datetime import datetime
from .base import BaseDocumentProcessor


class DOCXProcessor(BaseDocumentProcessor):
    """
    Microsoft Word (.docx) ファイルを処理するためのプロセッサークラス
    """

    def __init__(self):
        """
        DOCXプロセッサーの初期化
        """
        super().__init__()
        self.document = None

    def load(self, file_path: str) -> None:
        """
        DOCXファイルを読み込み、python-docxを使って解析します

        Args:
            file_path (str): DOCXファイルのパス
        """
        self.document = docx.Document(file_path)

    def extract_text(self) -> str:
        """
        DOCXからすべてのテキストを抽出します

        Returns:
            str: 抽出されたテキスト
        """
        if not self.document:
            return ""

        full_text = []

        # パラグラフからテキストを抽出
        for para in self.document.paragraphs:
            full_text.append(para.text)

        # テーブルからテキストを抽出
        for table in self.document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(' | '.join(row_text))

        return '\n\n'.join(full_text)

    def extract_metadata(self) -> Dict[str, Any]:
        """
        DOCXファイルからメタデータを抽出します

        Returns:
            Dict[str, Any]: メタデータのキーと値のペア
        """
        if not self.document:
            return {}

        metadata = {}

        # コアプロパティを取得
        properties = self.document.core_properties

        if properties:
            # 一般的なメタデータを抽出
            metadata.update({
                "author": properties.author,
                "title": properties.title,
                "subject": properties.subject,
                "keywords": properties.keywords,
                "category": properties.category,
                "comments": properties.comments,
                "created": properties.created.isoformat() if properties.created else None,
                "modified": properties.modified.isoformat() if properties.modified else None,
                "last_modified_by": properties.last_modified_by,
                "revision": properties.revision,
            })

        # 追加のドキュメント統計
        metadata.update({
            "paragraph_count": len(self.document.paragraphs),
            "table_count": len(self.document.tables),
            "section_count": len(self.document.sections),
        })

        # 文字数をカウント
        char_count = sum(len(p.text) for p in self.document.paragraphs)
        metadata["character_count"] = char_count

        return {k: v for k, v in metadata.items() if v is not None}

    def extract_sections(self) -> List[Dict[str, str]]:
        """
        DOCXからセクションを抽出します
        見出しスタイルを使用して文書を分割します

        Returns:
            List[Dict[str, str]]: セクションのリスト（タイトルとコンテンツを含む）
        """
        if not self.document:
            return super().extract_sections()

        sections = []
        current_heading = "はじめに"
        current_content = []

        for paragraph in self.document.paragraphs:
            # 見出しスタイルかどうかを確認
            if paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('見出し'):
                # 前のセクションを保存
                if current_content:
                    sections.append({
                        "title": current_heading,
                        "content": '\n'.join(current_content).strip()
                    })

                current_heading = paragraph.text
                current_content = []
            else:
                # 空でないパラグラフだけを追加
                if paragraph.text.strip():
                    current_content.append(paragraph.text)

        # 最後のセクションを追加
        if current_content:
            sections.append({
                "title": current_heading,
                "content": '\n'.join(current_content).strip()
            })

        # セクションが識別されなかった場合、デフォルトのセクション分けを使用
        if not sections:
            sections = super().extract_sections()

        return sections
